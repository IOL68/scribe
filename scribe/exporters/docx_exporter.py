"""
Exportador a formato DOCX (Word) - Estilo LAI
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_horizontal_line(paragraph, thick=False, color="808080"):
    """Agrega una línea horizontal debajo del párrafo."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12' if thick else '6')  # Grosor: 12=gruesa, 6=normal
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)  # Gris
    pBdr.append(bottom)
    pPr.append(pBdr)


def format_time(seconds: float) -> str:
    """Convierte segundos a formato MM:SS"""
    minutes = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{minutes:02d}:{secs:02d}"


def export_docx(transcription: dict, output_path: Path) -> None:
    """
    Exporta la transcripción a formato DOCX estilo LAI.

    Formato:
    - Interviewer:/Interviewee: como etiquetas
    - Hanging indent (texto nunca pasa debajo del label)
    - Times New Roman 12pt
    - [End of Audio] al final

    Args:
        transcription: Datos de la transcripción
        output_path: Ruta del archivo de salida
    """
    doc = Document()

    # Forzar fondo blanco (evita que Word use dark mode)
    shading_elm = OxmlElement('w:background')
    shading_elm.set(qn('w:color'), 'FFFFFF')
    doc.element.insert(0, shading_elm)

    # Configurar márgenes (estilo LAI)
    for section in doc.sections:
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Header: título del archivo + tipo de documento + dos líneas horizontales
    filename = output_path.stem if output_path else "Transcript"

    # Línea 1: Nombre del archivo (bold italic)
    header_para1 = doc.add_paragraph()
    header_para1.paragraph_format.space_after = Pt(0)
    header_run1 = header_para1.add_run(filename)
    header_run1.font.name = "Times New Roman"
    header_run1.font.size = Pt(12)
    header_run1.bold = True
    header_run1.italic = True

    # Línea 2: Interviewer/Interviewee (italic) + línea normal
    header_para2 = doc.add_paragraph()
    header_para2.paragraph_format.space_after = Pt(0)
    header_para2.paragraph_format.space_before = Pt(0)
    header_run2 = header_para2.add_run("Interviewer/Interviewee")
    header_run2.font.name = "Times New Roman"
    header_run2.font.size = Pt(12)
    header_run2.italic = True
    add_horizontal_line(header_para2, thick=False)  # Línea normal arriba

    # Línea gruesa (párrafo vacío con línea gruesa) - pegada a la anterior
    header_para3 = doc.add_paragraph()
    header_para3.paragraph_format.space_before = Pt(0)
    header_para3.paragraph_format.space_after = Pt(12)  # Espacio después de las líneas
    header_para3.paragraph_format.line_spacing = Pt(2)  # Línea muy pequeña
    add_horizontal_line(header_para3, thick=True)  # Línea gruesa abajo

    # Mapeo de speakers a roles
    # Por defecto: primer speaker = Interviewer, segundo = Interviewee
    speaker_map = {}
    first_speaker = None

    for segment in transcription["segments"]:
        speaker = segment.get("speaker", "Unknown")
        if speaker not in speaker_map and speaker != "Unknown":
            if first_speaker is None:
                first_speaker = speaker
                speaker_map[speaker] = "Interviewer"
            else:
                speaker_map[speaker] = "Interviewee"

    # Si solo hay un speaker
    if len(speaker_map) == 1:
        speaker_map[first_speaker] = "Speaker"

    # Agrupar segmentos por speaker
    current_speaker = None
    current_texts = []

    def flush_paragraph():
        """Escribe el párrafo acumulado al documento con hanging indent."""
        nonlocal current_texts, current_speaker

        if not current_texts or current_speaker is None:
            return

        role = speaker_map.get(current_speaker, current_speaker)

        # Crear párrafo
        para = doc.add_paragraph()

        # Configurar hanging indent (indentación colgante)
        # left=2160 twips (1.5"), hanging=-2160 twips
        para_format = para.paragraph_format
        para_format.left_indent = Twips(2160)  # 1.5 pulgadas
        para_format.first_line_indent = Twips(-2160)  # Negativo = hanging

        # Etiqueta (Interviewer: o Interviewee:) con tab - EN ITALIC
        label_run = para.add_run(f"{role}:\t")
        label_run.font.name = "Times New Roman"
        label_run.font.size = Pt(12)
        label_run.italic = True

        # Texto combinado - NORMAL (sin italic)
        combined_text = " ".join(current_texts)
        text_run = para.add_run(combined_text)
        text_run.font.name = "Times New Roman"
        text_run.font.size = Pt(12)
        text_run.italic = False

        # Reset
        current_texts = []

    for segment in transcription["segments"]:
        speaker = segment.get("speaker", "Unknown")
        text = segment.get("text", "").strip()
        timestamp = format_time(segment["start"])
        min_confidence = segment.get("min_confidence", 1.0)

        # Limpiar marcadores de Whisper (>>, --, etc.)
        text = text.lstrip(">").lstrip("-").lstrip().strip()
        if text.startswith(">>"):
            text = text[2:].strip()

        # Si cambia el speaker, flush el párrafo anterior
        if speaker != current_speaker:
            flush_paragraph()
            current_speaker = speaker

            # Agregar línea en blanco entre speakers
            if len(doc.paragraphs) > 0:
                empty_para = doc.add_paragraph()
                empty_para.paragraph_format.space_after = Pt(0)

        # Procesar texto - solo marcar si confianza muy baja
        processed_text = text

        # Solo marcar como unintelligible si confianza muy baja
        if min_confidence < 0.2:
            processed_text = f"[unintelligible {timestamp}]"

        current_texts.append(processed_text)

    # Flush último párrafo
    flush_paragraph()

    # Agregar [End of Audio] - en italic
    doc.add_paragraph()
    end_para = doc.add_paragraph()
    end_run = end_para.add_run("[End of Audio]")
    end_run.font.name = "Times New Roman"
    end_run.font.size = Pt(12)
    end_run.italic = True

    # Footer con dos líneas horizontales + www.thelai.com y número de página
    for section in doc.sections:
        footer = section.footer

        # Línea gruesa arriba
        line_para1 = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        line_para1.clear()
        line_para1.paragraph_format.space_after = Pt(0)
        line_para1.paragraph_format.line_spacing = Pt(2)
        add_horizontal_line(line_para1, thick=True)

        # Línea normal - pegada
        line_para2 = footer.add_paragraph()
        line_para2.paragraph_format.space_before = Pt(0)
        line_para2.paragraph_format.space_after = Pt(0)
        line_para2.paragraph_format.line_spacing = Pt(2)
        add_horizontal_line(line_para2, thick=False)

        # Párrafo del footer con contenido (en la misma línea)
        footer_para = footer.add_paragraph()
        footer_para.paragraph_format.space_before = Pt(0)

        # Tab stop alineado a la derecha (5.5" considerando márgenes de 1.25")
        tab_stops = footer_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT)

        # www.thelai.com (izquierda)
        run1 = footer_para.add_run("www.thelai.com")
        run1.font.name = "Arial"
        run1.font.size = Pt(9)
        run1.italic = True

        # Tab para ir a la derecha
        footer_para.add_run("\t")

        # Page X of Y
        run2 = footer_para.add_run("Page ")
        run2.font.name = "Arial"
        run2.font.size = Pt(9)
        run2.italic = True

        # Campo PAGE (número de página actual)
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.text = "PAGE"
        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'end')

        run_page = footer_para.add_run()
        run_page._r.append(fld_char1)
        run_page._r.append(instr)
        run_page._r.append(fld_char2)
        run_page.font.name = "Arial"
        run_page.font.size = Pt(9)
        run_page.italic = True

        run3 = footer_para.add_run(" of ")
        run3.font.name = "Arial"
        run3.font.size = Pt(9)
        run3.italic = True

        # Campo NUMPAGES (total de páginas)
        fld_char3 = OxmlElement('w:fldChar')
        fld_char3.set(qn('w:fldCharType'), 'begin')
        instr2 = OxmlElement('w:instrText')
        instr2.text = "NUMPAGES"
        fld_char4 = OxmlElement('w:fldChar')
        fld_char4.set(qn('w:fldCharType'), 'end')

        run_numpages = footer_para.add_run()
        run_numpages._r.append(fld_char3)
        run_numpages._r.append(instr2)
        run_numpages._r.append(fld_char4)
        run_numpages.font.name = "Arial"
        run_numpages.font.size = Pt(9)
        run_numpages.italic = True

    # Guardar
    doc.save(output_path)
